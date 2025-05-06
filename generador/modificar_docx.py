# generador/modificar_docx.py
from docx import Document
import re
from pathlib import Path
from typing import Dict

def _build_regex_mapping(replacements: Dict[str, str]) -> Dict[re.Pattern, str]:
    """Convierte {texto_original: texto_nuevo} en {regex: texto_nuevo} ignorando mayúsculas/minúsculas."""
    return {re.compile(re.escape(k), re.IGNORECASE): v for k, v in replacements.items()}

def reemplazar_campos_docx(plantilla: Path, salida: Path, reemplazos: Dict[str, str]) -> None:
    """
    Crea una copia de `plantilla` en `salida` y reemplaza cada literal (clave)
    por su valor correspondiente en todo el documento (párrafos + tablas).
    """
    doc = Document(plantilla)
    regex_map = _build_regex_mapping(reemplazos)

    def _reemplazar_en_runs(runs):
        for run in runs:
            text = run.text
            for patron, nuevo in regex_map.items():
                if patron.search(text):
                    run.text = patron.sub(str(nuevo), text)

    # Párrafos normales
    for p in doc.paragraphs:
        _reemplazar_en_runs(p.runs)

    # Párrafos dentro de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _reemplazar_en_runs(p.runs)

    doc.save(salida)
    print(f"✅ Documento generado en: {salida}")
