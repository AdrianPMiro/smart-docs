# backend/services/docx_template_detector.py

from pathlib import Path
from docx import Document  # python-docx
from typing import Optional

# --- Lista de identificadores únicos para cada tipo de PLANTILLA DOCX ---
# (texto_unico_a_buscar, tipo_a_devolver)
# Estos textos deben ser muy específicos y estar presentes solo en ese tipo de plantilla.
# ¡DEBES AJUSTAR ESTOS IDENTIFICADORES!
DOCX_TEMPLATE_IDENTIFIERS = [
    # Ejemplo para una plantilla de contrato de alquiler
    ("CONTRATO DE ARRENDAMIENTO DE VIVIENDA", "plantilla_contrato_alquiler_v1"),
    ("CLÁUSULAS ESPECÍFICAS DEL ARRENDADOR", "plantilla_contrato_alquiler_v1"),

    # Ejemplo para una plantilla de carta de oferta laboral
    ("CARTA DE OFERTA LABORAL", "plantilla_oferta_empleo_estandar"),
    ("Condiciones de la Incorporación", "plantilla_oferta_empleo_estandar"),

    # Ejemplo para un boletín/certificado (si tus plantillas DOCX son así)
    # Asumiendo que el "boletin.xlsx" que mencionaste antes tiene un equivalente DOCX
    ("CERTIFICADO DE INSTALACIÓN DE BAJA TENSIÓN", "plantilla_certificado_bt_docx"),

    # Añade más tuplas para cada tipo de plantilla DOCX
]


def detect_docx_template_type(docx_path: Path) -> Optional[str]:
    """
    Detecta el tipo de una plantilla DOCX basado en cadenas de texto únicas.

    Args:
        docx_path (Path): Ruta al archivo DOCX.

    Returns:
        Optional[str]: El string identificador del tipo de plantilla DOCX, o None si no se reconoce.
    """
    if not docx_path.exists() or docx_path.suffix.lower() != ".docx":
        return None

    try:
        doc = Document(docx_path)
        # Leer texto de los primeros N párrafos (ej. los primeros 10-20)
        # o un número limitado de caracteres para eficiencia.
        text_to_search = ""
        # Leer algunos párrafos del inicio del documento
        for i, para in enumerate(doc.paragraphs):
            if i >= 20:  # Limitar a los primeros 20 párrafos
                break
            text_to_search += para.text + "\n"

        # Opcionalmente, también podrías revisar texto en tablas si es relevante
        # for table in doc.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             text_to_search += cell.text + "\n"
        #             if len(text_to_search) > 5000: break # Límite de caracteres
        #     if len(text_to_search) > 5000: break
        # if len(text_to_search) > 5000: text_to_search = text_to_search[:5000]

        for identifier_text, template_type in DOCX_TEMPLATE_IDENTIFIERS:
            # Búsqueda sensible a mayúsculas por defecto
            if identifier_text in text_to_search:
                return template_type

    except Exception as e:
        print(f"Error al procesar DOCX {docx_path} para detección de tipo: {e}")
        return None

    return "docx_plantilla_desconocida"  # Tipo por defecto