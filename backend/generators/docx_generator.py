import re
from pathlib import Path
from typing import Dict
from docx import Document

try:
    from backend.readers.docx_reader import docx_reader
except ImportError:
    print("ADVERTENCIA: No se pudo importar 'docx_reader' desde 'readers'.")
    def docx_reader(path: Path) -> Dict[str, str]:
        print(f"Llamada a docx_reader (dummy) para {path}. Devolviendo diccionario vacío.")
        return {}

def _build_regex_mapping(replacements_map: Dict[str, str]) -> Dict[re.Pattern, str]:
    return {
        re.compile(re.escape(k), re.IGNORECASE): str(v)
        for k, v in replacements_map.items() if k
    }

def _reemplazar_texto_en_runs(runs_list: list, regex_map_compiled: Dict[re.Pattern, str]):
    for run_item in runs_list:
        original_text = run_item.text
        modified_text = original_text
        for pattern, new_val in regex_map_compiled.items():
            modified_text = pattern.sub(new_val, modified_text)
        if modified_text != original_text:
            run_item.text = modified_text

def _aplicar_reemplazos_al_documento(doc: Document, regex_map_compiled: Dict[re.Pattern, str]):
    for paragraph_item in doc.paragraphs:
        _reemplazar_texto_en_runs(paragraph_item.runs, regex_map_compiled)
    for table_item in doc.tables:
        for row_item in table_item.rows:
            for cell_item in row_item.cells:
                for paragraph_item_in_cell in cell_item.paragraphs:
                    _reemplazar_texto_en_runs(paragraph_item_in_cell.runs, regex_map_compiled)

def docx_generator(
        plantilla_path: Path,
        output_path: Path,
        datos_nuevos: Dict[str, str]
) -> None:
    print(f" [DOCX Generator] Iniciando generación para: {output_path}")
    print(f" [DOCX Generator] Usando plantilla: {plantilla_path}")
    print(f" [DOCX Generator] Extrayendo datos actuales (datos2) de la plantilla...")
    datos_plantilla_actuales = docx_reader(plantilla_path)
    if not datos_plantilla_actuales:
        print(f"⚠️  [DOCX Generator] No se pudieron extraer datos de la plantilla {plantilla_path}.")
        print(f"⚠️  [DOCX Generator] Se guardará una copia de la plantilla sin modificaciones en {output_path}.")
        try:
            doc_copy = Document(plantilla_path)
            doc_copy.save(output_path)
        except Exception as e:
            print(f"❌ [DOCX Generator] Error al guardar copia de la plantilla: {e}")
        return
    mapa_de_reemplazos: Dict[str, str] = {}
    campos_comunes = [
        'nif', 'nombre', 'apellido1', 'apellido2', 'email',
        'direccion', 'cp', 'localidad', 'provincia', 'telefono',
        'tipo_via', 'nom_via', 'numero'
    ]
    for campo in campos_comunes:
        valor_placeholder_en_plantilla = datos_plantilla_actuales.get(campo, "")
        valor_nuevo_de_datos = datos_nuevos.get(campo, "")
        if valor_placeholder_en_plantilla:
            mapa_de_reemplazos[valor_placeholder_en_plantilla] = valor_nuevo_de_datos
    if not mapa_de_reemplazos:
        print("⚠️  [DOCX Generator] No se generó ningún mapeo de reemplazo válido.")
        print(f"⚠️  [DOCX Generator] Se guardará una copia de la plantilla sin modificaciones en {output_path}.")
        try:
            doc_copy_no_map = Document(plantilla_path)
            doc_copy_no_map.save(output_path)
        except Exception as e:
            print(f"❌ [DOCX Generator] Error al guardar copia de la plantilla (sin mapa): {e}")
        return
    try:
        doc_a_modificar = Document(plantilla_path)
    except Exception as e:
        print(f"❌ [DOCX Generator] Error al abrir la plantilla DOCX para reemplazo: {plantilla_path}. Detalle: {e}")
        return
    regex_map_compilado = _build_regex_mapping(mapa_de_reemplazos)
    if not regex_map_compilado:
        print("⚠️  [DOCX Generator] El mapa de expresiones regulares compilado está vacío.")
        print(f"⚠️  [DOCX Generator] Se guardará una copia de la plantilla sin modificaciones en {output_path}.")
        try:
            doc_a_modificar.save(output_path)
        except Exception as e:
            print(f"❌ [DOCX Generator] Error al guardar copia de la plantilla (regex map vacío): {e}")
        return
    _aplicar_reemplazos_al_documento(doc_a_modificar, regex_map_compilado)
    try:
        doc_a_modificar.save(output_path)
        print(f"[DOCX Generator] Documento generado exitosamente en: {output_path}")
    except Exception as e:
        print(f"❌ [DOCX Generator] Error al guardar el documento DOCX modificado: {output_path}. Detalle: {e}")